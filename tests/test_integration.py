"""End-to-end integration tests for AllPDF.

Requires: LibreOffice to be installed for Office->PDF tests.
"""
import os
import subprocess
import pytest
from pathlib import Path
from allpdf.orchestrator import Orchestrator
from allpdf.models import FileFormat


def _has_libreoffice():
    try:
        result = subprocess.run(["soffice", "--version"], capture_output=True, timeout=10)
        return result.returncode == 0
    except Exception:
        return False


requires_libreoffice = pytest.mark.skipif(
    not _has_libreoffice(), reason="LibreOffice is not installed or not on PATH"
)


class TestDocxToPdf:
    @requires_libreoffice
    def test_simple_docx_to_pdf(self, tmp_path):
        from docx import Document
        docx_path = tmp_path / "hello.docx"
        d = Document()
        d.add_heading("AllPDF Integration Test", level=1)
        d.add_paragraph("This is a paragraph with some text.")
        d.add_paragraph("Another paragraph for testing.")
        d.save(str(docx_path))

        pdf_path = tmp_path / "hello.pdf"
        orch = Orchestrator()
        result = orch.convert(
            str(docx_path), str(pdf_path),
            FileFormat.DOCX, FileFormat.PDF,
            quality_check=True,
        )
        assert result.status.value == "success"
        assert os.path.exists(str(pdf_path))
        assert result.quality_report is not None
        assert result.quality_report.overall_grade.value != "red"

    @requires_libreoffice
    def test_docx_with_table_to_pdf(self, tmp_path):
        from docx import Document
        docx_path = tmp_path / "table.docx"
        d = Document()
        d.add_heading("Table Test", level=1)
        table = d.add_table(rows=3, cols=3)
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.text = f"R{i}C{j}"
        d.save(str(docx_path))

        pdf_path = tmp_path / "table.pdf"
        orch = Orchestrator()
        result = orch.convert(
            str(docx_path), str(pdf_path),
            FileFormat.DOCX, FileFormat.PDF,
        )
        assert result.status.value == "success"
        assert os.path.exists(str(pdf_path))


class TestPdfToDocx:
    def test_simple_pdf_to_docx(self, tmp_path):
        import fitz
        pdf_path = tmp_path / "simple.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Integration test PDF content", fontsize=12)
        page.insert_text((72, 100), "Second line of text", fontsize=12)
        doc.save(str(pdf_path))
        doc.close()

        docx_path = tmp_path / "simple.docx"
        orch = Orchestrator()
        result = orch.convert(
            str(pdf_path), str(docx_path),
            FileFormat.PDF, FileFormat.DOCX,
            quality_check=True,
        )
        assert result.status.value == "success"
        assert os.path.exists(str(docx_path))
        assert os.path.getsize(str(docx_path)) > 0

    def test_multipage_pdf_to_docx(self, tmp_path):
        import fitz
        pdf_path = tmp_path / "multi.pdf"
        doc = fitz.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((72, 72), f"Page {i+1}", fontsize=14)
        doc.save(str(pdf_path))
        doc.close()

        docx_path = tmp_path / "multi.docx"
        orch = Orchestrator()
        result = orch.convert(
            str(pdf_path), str(docx_path),
            FileFormat.PDF, FileFormat.DOCX,
        )
        assert result.status.value == "success"


class TestErrorHandling:
    def test_nonexistent_file(self, tmp_path):
        orch = Orchestrator()
        result = orch.convert(
            str(tmp_path / "ghost.pdf"),
            str(tmp_path / "ghost.docx"),
            FileFormat.PDF, FileFormat.DOCX,
        )
        assert result.status.value == "failed"
        assert result.error_message is not None

    def test_unsupported_format_pair(self, tmp_path):
        orch = Orchestrator()
        result = orch.convert(
            str(tmp_path / "a.xlsx"),
            str(tmp_path / "a.pptx"),
            FileFormat.XLSX, FileFormat.PPTX,
        )
        assert result.status.value == "failed"
        assert "No engine" in result.error_message

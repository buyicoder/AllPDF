# tests/test_quality/test_checker.py
import os
from pathlib import Path
import fitz
from allpdf.quality.checker import QualityChecker
from allpdf.models import FileFormat, QualityGrade, FileInfo


def _make_pdf(path: str, pages: int = 2, text: str = "Hello"):
    doc = fitz.open()
    for i in range(pages):
        page = doc.new_page()
        page.insert_text((72, 72), f"{text} page {i+1}", fontsize=12)
    doc.save(path)
    doc.close()


class TestQualityChecker:
    def test_check_pages_match(self, tmp_path):
        pdf = tmp_path / "test.pdf"
        _make_pdf(str(pdf), pages=3)

        source_info = FileInfo(
            path=pdf, format=FileFormat.PDF,
            page_count=3, image_count=0, table_count=0,
            is_scanned=False, file_size_bytes=100,
        )
        result_info = FileInfo(
            path=Path("/tmp/out.docx"), format=FileFormat.DOCX,
            page_count=3, image_count=0, table_count=0,
            is_scanned=False, file_size_bytes=200,
        )

        checker = QualityChecker()
        check = checker.check_page_count(source_info, result_info, FileFormat.PDF, FileFormat.DOCX)
        assert check.grade == QualityGrade.GREEN

    def test_check_pages_mismatch_warning(self):
        source_info = FileInfo(
            path=Path("in.pdf"), format=FileFormat.PDF,
            page_count=5, image_count=0, table_count=0,
            is_scanned=False, file_size_bytes=100,
        )
        result_info = FileInfo(
            path=Path("out.docx"), format=FileFormat.DOCX,
            page_count=3, image_count=0, table_count=0,
            is_scanned=False, file_size_bytes=200,
        )

        checker = QualityChecker()
        check = checker.check_page_count(source_info, result_info, FileFormat.PDF, FileFormat.DOCX)
        assert check.grade == QualityGrade.YELLOW

    def test_check_file_openable(self, tmp_path):
        docx_path = tmp_path / "test.docx"
        from docx import Document
        d = Document()
        d.add_paragraph("test")
        d.save(str(docx_path))

        checker = QualityChecker()
        check = checker.check_openable(str(docx_path), FileFormat.DOCX)
        assert check.grade == QualityGrade.GREEN

    def test_check_file_not_openable(self, tmp_path):
        bad_path = tmp_path / "corrupt.docx"
        bad_path.write_text("not a real docx file")

        checker = QualityChecker()
        check = checker.check_openable(str(bad_path), FileFormat.DOCX)
        assert check.grade == QualityGrade.RED

    def test_full_check_all_green(self, tmp_path):
        pdf = tmp_path / "src.pdf"
        _make_pdf(str(pdf), pages=1)

        docx = tmp_path / "out.docx"
        from docx import Document
        d = Document()
        d.add_paragraph("Hello page 1")
        d.save(str(docx))

        checker = QualityChecker()
        report = checker.run(
            input_path=str(pdf),
            output_path=str(docx),
            input_format=FileFormat.PDF,
            output_format=FileFormat.DOCX,
        )
        assert report.overall_grade == QualityGrade.GREEN
